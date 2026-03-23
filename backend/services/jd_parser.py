import os
import io
import pdfplumber
from docx import Document


def parse_jd_from_bytes(file_bytes: bytes, filename: str) -> str:
    """Extract text directly from file bytes — no temp file needed."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        text = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()

    elif ext in (".docx", ".doc"):
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        text = "\n".join(paragraphs)

        # Also check tables in case JD content is inside a table
        if not text.strip():
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
        return text.strip()

    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload PDF or DOCX.")